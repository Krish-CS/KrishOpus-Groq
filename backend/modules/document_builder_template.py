"""
Document Builder V24.0 - ULTIMATE ROBUST
✅ Handles "Assignment Topic" OR "Title of Journal Paper"
✅ Handles with/without COURSE TEACHER SIGNATURE
✅ Handles with/without placeholder text
✅ Handles templates with existing content
✅ Smart topic replacement
✅ Smart content deletion after signatures
✅ Proper page break
"""

import os
import re
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class SmartDocumentBuilder:
    """V24 - ULTIMATE ROBUST"""
    
    def __init__(self, ollama_url: str = "http://localhost:11434"):
        self.ollama_url = ollama_url
        print("✓ Document Builder initialized (V24 - ULTIMATE ROBUST)")
    
    
    def analyze_template(self, template_path: str) -> List[str]:
        """Analyze template for sections"""
        try:
            doc = Document(template_path)
            
            for table in doc.tables:
                first_row = ' '.join([cell.text for cell in table.rows[0].cells]).lower()
                
                if ('objective' in first_row or 'descriptive' in first_row) and 'conclusion' in first_row:
                    sections = []
                    for cell in table.rows[0].cells:
                        text = cell.text.strip().replace('\n', ' ').split('(')[0].strip()
                        if not text or len(text) < 2:
                            continue
                        text_lower = text.lower()
                        if text_lower in ['total', 'marks', 'awarded']:
                            continue
                        if 'conclusion' in text_lower and 'reference' in text_lower:
                            sections.append('Conclusion')
                            continue
                        if text_lower in ['reference', 'references']:
                            continue
                        sections.append(text)
                    sections.append('References')
                    return sections
            
            return ['Objective', 'Context and Relevance', 'Technology aspects and relevant drawings', 'Emerging trends ideas you visualize in the field', 'Conclusion', 'References']
        except:
            return ['Objective', 'Context and Relevance', 'Technology aspects and relevant drawings', 'Emerging trends ideas you visualize in the field', 'Conclusion', 'References']
    
    
    def generate_from_template(
        self,
        template_path: str,
        topic: str,
        user_data: Dict,
        generated_content: Dict,
        output_format: str = "docx"
    ) -> Dict:
        """Generate assignment - ULTIMATE ROBUST"""
        
        result = {'status': 'success', 'output_path': None}
        
        try:
            doc = Document(template_path)
            
            print(f"\n🔧 Step 1: Smart topic replacement...")
            topic_filled = self._fill_topic_smart(doc, topic)
            if topic_filled:
                print(f"   ✅ Topic filled: {topic}")
            else:
                print(f"   ⚠ Topic field not found")
            
            print(f"🗑️ Step 2: Smart content deletion...")
            deleted = self._delete_content_smart(doc)
            print(f"   ✅ Deleted {deleted} elements")
            
            print(f"📄 Step 3: Adding page break...")
            self._add_page_break_smart(doc)
            
            print(f"📝 Step 4: Adding our content on page 2...")
            self._add_content_page2(doc, topic, generated_content)
            
            # Save
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_name = "".join(c for c in user_data.get('name', 'Student') if c.isalnum())[:15]
            if not safe_name:
                safe_name = "Student"
            output_filename = f"Assignment_{safe_name}_{timestamp}.docx"
            
            output_dir = "outputs"
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, output_filename)
            
            doc.save(output_path)
            
            result['output_path'] = output_path
            print(f"✅ Saved: {output_filename}\n")
            
            return result
            
        except Exception as e:
            print(f"❌ Error: {e}")
            import traceback
            traceback.print_exc()
            result['status'] = 'error'
            result['message'] = str(e)
            return result
    
    
    def _fill_topic_smart(self, doc: Document, topic: str) -> bool:
        """
        ✅ Smart topic replacement - handles multiple formats:
        - "Assignment Topic"
        - "Title of Journal Paper"
        - Existing topic text (replaces it)
        """
        
        print(f"   Looking for topic field...")
        
        # Keywords to identify topic rows
        topic_keywords = ['Assignment Topic', 'Title of Journal Paper', 'Journal Paper']
        
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    
                    # Check if this cell contains topic keyword
                    for keyword in topic_keywords:
                        if keyword in cell_text:
                            # Found topic row! Next row contains the topic cell
                            if row_idx + 1 < len(table.rows):
                                topic_row = table.rows[row_idx + 1]
                                topic_cell = topic_row.cells[0]
                                
                                # Clear the cell
                                for para in topic_cell.paragraphs:
                                    for run in para.runs:
                                        run.text = ''
                                
                                # Add topic
                                if topic_cell.paragraphs:
                                    para = topic_cell.paragraphs[0]
                                else:
                                    para = topic_cell.add_paragraph()
                                
                                para.clear()
                                font_size = 9 if len(topic) > 60 else 10
                                
                                run = para.add_run(topic)
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(font_size)
                                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                
                                print(f"   ✓ Found '{keyword}' in table {table_idx}, row {row_idx}")
                                return True
        
        return False
    
    
    def _delete_content_smart(self, doc: Document) -> int:
        """
        ✅ Smart content deletion:
        1. If COURSE TEACHER SIGNATURE found: Delete everything AFTER it
        2. If "Type Your Heading Here" found: Delete from there onwards
        3. If numbered list (1. 2. 3...) found: Delete from there onwards
        4. Otherwise: Don't delete anything (clean template)
        """
        
        try:
            paragraphs = list(doc.paragraphs)
            deletion_start_idx = None
            deletion_reason = None
            
            # Find deletion starting point
            for idx, para in enumerate(paragraphs):
                para_text = para.text.strip()
                
                # Strategy 1: Delete after COURSE TEACHER SIGNATURE
                if 'COURSE TEACHER' in para_text.upper() or 'SIGNATURE' in para_text.upper():
                    # Look for next 3 paragraphs for instructor name
                    for j in range(idx + 1, min(idx + 4, len(paragraphs))):
                        next_text = paragraphs[j].text.strip()
                        if 'Mrs.' in next_text or 'Mr.' in next_text or 'Dr.' in next_text:
                            # Delete AFTER instructor name
                            deletion_start_idx = j + 1
                            deletion_reason = f"after COURSE TEACHER SIGNATURE ({next_text})"
                            break
                    
                    if deletion_start_idx:
                        break
                
                # Strategy 2: Delete from "Type Your Heading Here"
                if 'Type Your Heading' in para_text or 'type your heading' in para_text.lower():
                    deletion_start_idx = idx
                    deletion_reason = "'Type Your Heading Here' placeholder"
                    break
                
                # Strategy 3: Delete numbered list (but not if it's in first 10 paragraphs - those are headers)
                if idx > 10 and re.match(r'^\d+\.\s+\w+', para_text):
                    if deletion_start_idx is None:
                        deletion_start_idx = idx
                        deletion_reason = f"numbered list starting with '{para_text[:30]}'"
            
            # Perform deletion
            if deletion_start_idx is not None:
                paragraphs_to_delete = paragraphs[deletion_start_idx:]
                for para in paragraphs_to_delete:
                    p = para._element
                    p.getparent().remove(p)
                
                print(f"   ℹ Deleted from {deletion_reason}")
                return len(paragraphs_to_delete)
            else:
                print(f"   ℹ No placeholder content found (clean template)")
                return 0
            
        except Exception as e:
            print(f"   ⚠ Deletion error: {e}")
            return 0
    
    
    def _add_page_break_smart(self, doc: Document):
        """Add proper page break"""
        
        try:
            paragraphs = doc.paragraphs
            if paragraphs:
                # Find last non-empty paragraph
                last_para = None
                for para in reversed(paragraphs):
                    if para.text.strip():
                        last_para = para
                        break
                
                if last_para is None:
                    last_para = paragraphs[-1] if paragraphs else None
                
                if last_para:
                    # Add page break
                    run = last_para.add_run()
                    run.add_break(type=6)  # Page break
                    print(f"   ✓ Page break added")
                else:
                    # Fallback
                    doc.add_page_break()
                    print(f"   ✓ Page break added (fallback)")
            else:
                doc.add_page_break()
                print(f"   ✓ Page break added (no paragraphs)")
                
        except Exception as e:
            print(f"   ⚠ Page break error: {e}")
            doc.add_page_break()
            print(f"   ✓ Page break added (error fallback)")
    
    
    def _add_content_page2(self, doc: Document, topic: str, generated_content: Dict):
        """Add our content on page 2"""
        
        # Topic heading
        topic_para = doc.add_paragraph()
        run = topic_para.add_run(topic)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        topic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        topic_para.paragraph_format.space_before = Pt(0)
        topic_para.paragraph_format.space_after = Pt(12)
        
        print(f"   ✓ Topic heading")
        
        # Sections
        section_num = 1
        total_sections = len([k for k in generated_content.keys() if generated_content[k].strip()])
        
        for section_name, content in generated_content.items():
            if not content or len(content.strip()) < 10:
                continue
            
            # Section heading
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(f"{section_num}. {section_name}")
            heading_run.font.name = 'Times New Roman'
            heading_run.font.size = Pt(12)
            heading_run.font.bold = True
            heading_para.paragraph_format.space_before = Pt(0)
            heading_para.paragraph_format.space_after = Pt(6)
            
            # Content
            if 'reference' in section_name.lower():
                self._add_references(doc, content)
            else:
                para = doc.add_paragraph(content.strip())
                self._format_para(para)
            
            # Spacing between sections
            if section_num < total_sections:
                doc.add_paragraph()
            
            print(f"   ✓ Section {section_num}: {section_name}")
            section_num += 1
        
        print(f"   ✅ Added {section_num - 1} sections")
    
    
    def _format_para(self, para):
        """Format paragraph"""
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing = 1.15
    
    
    def _add_references(self, doc: Document, content: str):
        """Add numbered references"""
        refs = [line.strip() for line in content.split('\n') if line.strip()]
        
        ref_num = 1
        for ref in refs:
            ref = re.sub(r'^[•→★▸➤◆■\-●○]\s*', '', ref)
            ref = re.sub(r'^\d+[\.\)\\]\s*', '', ref)
            
            para = doc.add_paragraph(f"{ref_num}. {ref}")
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            
            para.paragraph_format.space_after = Pt(4)
            
            ref_num += 1
        
        print(f"   ✓ {ref_num - 1} references")
    
    
    def build_document(self, content: dict, user_data: dict, template_path: str = None) -> str:
        """Compatibility wrapper"""
        print("\n🔄 Converting format...")
        
        topic = content.get('title', 'Assignment')
        generated_content = {}
        for section in content.get('sections', []):
            generated_content[section.get('title', 'Section')] = section.get('content', '')
        
        if not template_path or not os.path.exists(template_path):
            for tmpl in ["templates/test_template.docx", "test_template.docx", "uploads/test_template.docx"]:
                if os.path.exists(tmpl):
                    template_path = tmpl
                    break
        
        result = self.generate_from_template(template_path, topic, user_data, generated_content)
        
        if result.get('status') == 'success':
            return result.get('output_path', '')
        else:
            raise Exception(result.get('message', 'Failed'))
